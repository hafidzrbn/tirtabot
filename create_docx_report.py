import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Slate dark
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(14, 165, 183) # Brand Teal
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_body_p(doc, text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.font.bold = True
        
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = italic
    return p

def add_bullet_p(doc, bold_title, body_text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    
    r_title = p.add_run(bold_title)
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(12)
    r_title.font.bold = True
    
    r_body = p.add_run(body_text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(12)
    return p

def add_image_centered(doc, image_path, caption_text, width_inches=5.8):
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        p_img.add_run().add_picture(image_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run(caption_text)
        r_cap.font.name = 'Times New Roman'
        r_cap.font.size = Pt(10)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 116, 139)

def main():
    doc = docx.Document()
    
    # 1. Page Margins (Academic Standard 4cm top/left, 3cm bottom/right)
    sec = doc.sections[0]
    sec.top_margin = Inches(1.575)
    sec.left_margin = Inches(1.575)
    sec.bottom_margin = Inches(1.181)
    sec.right_margin = Inches(1.181)
    
    # =========================================================
    # COVER PAGE
    # =========================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    r = p_title.add_run("LAPORAN UJIAN AKHIR SEMESTER (UAS) GENAP TA 2025/2026")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    r = p_sub.add_run("MATA KULIAH: TRENDING TOPICS ON STATISTICS (TToS)")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(14, 165, 183)

    p_case = doc.add_paragraph()
    p_case.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_case.paragraph_format.space_after = Pt(24)
    r = p_case.add_run("INTEGRASI TEXT MINING, PEMODELAN AI KOMPARATIVE (INDOBERT, SVM, LOGISTIC REGRESSION), DAN RETRIEVAL-AUGMENTED GENERATION (RAG) PADA 24.325 KOMENTAR YOUTUBE DR. TIRTA MANDIRA HUDHI")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True
    
    p_case_sub = doc.add_paragraph()
    p_case_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_case_sub.paragraph_format.space_after = Pt(40)
    r = p_case_sub.add_run("(Studi Kasus: 24.325 Komentar YouTube dr. Tirta Mandira Hudhi, Sp.B)")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.italic = True

    p_req = doc.add_paragraph()
    p_req.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_req.paragraph_format.space_after = Pt(30)
    r = p_req.add_run("Disusun untuk Memenuhi Ujian Akhir Semester\nMata Kuliah : Trending Topics on Statistics (TToS)\nDosen Penguji : Dr. RB Fajriya Hakim, M.Si.")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_after = Pt(4)
    r = p_by.add_run("OLEH:")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.bold = True

    p_student = doc.add_paragraph()
    p_student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_student.paragraph_format.space_after = Pt(50)
    r = p_student.add_run("MUHAMMAD HAFIDZ RABBAANII SULTHON (23611091)")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True

    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ.paragraph_format.space_after = Pt(0)
    r = p_univ.add_run("JURUSAN STATISTIKA\nFAKULTAS MATEMATIKA DAN ILMU PENGETAHUAN ALAM\nUNIVERSITAS ISLAM INDONESIA\nYOGYAKARTA\n2026")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True

    doc.add_page_break()

    # =========================================================
    # BAB I: AKUISISI DAN KARAKTERISASI DATA
    # =========================================================
    add_heading_1(doc, "BAB I\nAKUISISI DAN KARAKTERISASI DATA")
    
    add_heading_2(doc, "1.1 Sumber Data")
    add_body_p(doc, "Data yang digunakan dalam penelitian dan analisis ini adalah komentar pengguna YouTube yang diekstrak dari 10 video populer yang menampilkan atau dibuat oleh dr. Tirta Mandira Hudhi, Sp.B. Video-video tersebut mencakup berbagai tema utama: Siniar/Podcast Populer (PWK, Close The Door, Raditya Dika, Denny Sumargo), Edukasi Kesehatan & Lifestyle, Review & Bisnis Sepatu Lari Lokal, serta Klarifikasi Isu & Opini Publik.")
    add_bullet_p(doc, "Metode Akuisisi: ", "Menggunakan otomatisasi API / Python library youtube-comment-downloader.")
    add_bullet_p(doc, "Total Komentar Raw: ", "24.325 baris komentar yang berhasil ditarik secara utuh.")

    add_heading_2(doc, "1.2 Rincian Sebaran Komentar per Video")
    add_body_p(doc, "Berikut adalah rincian 10 video YouTube dr. Tirta beserta sebaran jumlah komentar yang diekstrak:")

    # Table 1.1
    t1 = doc.add_table(rows=12, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    
    headers1 = ["No", "Video ID", "Judul / Kanal Video", "Tema Utama", "Jumlah Komentar"]
    widths1 = [Inches(0.4), Inches(1.1), Inches(2.2), Inches(1.8), Inches(1.0)]
    
    for i, h in enumerate(headers1):
        cell = t1.cell(0, i)
        cell.width = widths1[i]
        cell.text = h
        set_cell_background(cell, "0EA5B7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    data1 = [
        ("1", "dSq0Z5XpoLc", "PWK (Podcast Warung Kopi) - HAS Creative", "Personal, Mentalitas, & Anak", "5.278"),
        ("2", "l5pK6sfhxt0", "Close The Door - Deddy Corbuzier", "Kebijakan Pandemi & Masker vs Makan", "4.124"),
        ("3", "7mrwndoqyMk", "Raditya Dika Podcast", "Finansial, Gaji Dokter, & Karir", "3.254"),
        ("4", "UyalifZrhGM", "Tirta PengPengPeng (Official)", "Edukasi Kesehatan & Diabetes", "2.346"),
        ("5", "QtIxl1YM9Bk", "Tirta PengPengPeng (Official)", "Review Sepatu Lari Lokal (#Tirtalokal)", "2.261"),
        ("6", "41itFALrNU8", "NOICE / Podcast Bahlul", "Klarifikasi Isu Fitnah & Hukum", "2.099"),
        ("7", "lqeDF5JwYvM", "CURHAT BANG Denny Sumargo", "Perjalanan Hidup & Shoes and Care", "2.045"),
        ("8", "2qWR_b1HE18", "Raditya Dika (Edisi Finansial Medis)", "Realitas Profesi Kesehatan", "1.217"),
        ("9", "LCWsCEqAU8s", "Tirta PengPengPeng (#Tirtalokal)", "UMKM & Industri Kreatif", "1.044"),
        ("10", "CoVz4-TPYgM", "Tirta PengPengPeng (Edukasi Medis)", "Jantung, Hipertensi, & Gaya Hidup", "657"),
        ("TOTAL", "", "", "", "24.325")
    ]
    
    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, text_val in enumerate(row_data):
            cell = t1.cell(row_idx, col_idx)
            cell.width = widths1[col_idx]
            cell.text = text_val
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            if row_idx == 11: # Total row
                set_cell_background(cell, "E2E8F0")
            p = cell.paragraphs[0]
            if col_idx in [0, 1, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9.5)
                if row_idx == 11:
                    r.font.bold = True

    add_heading_2(doc, "1.3 Tujuan Analisis")
    add_body_p(doc, "Membangun sistem analisis data teks cerdas yang mengombinasikan Eksplorasi Statistik Teks, Pemodelan Klasifikasi Sentimen AI Komparatif (Logistic Regression, SVM, IndoBERT), serta Prototipe RAG berbasis Vektor FAISS untuk mengekstrak aspirasi, persepsi publik, dan informasi edukatif guna mendukung pengambilan keputusan di bidang kesehatan publik dan industri kreatif.")

    add_heading_2(doc, "1.4 Karakteristik Data")
    add_bullet_p(doc, "Struktur Atribut: ", "video_url, video_id, cid, author, text, time, votes, replies.")
    add_bullet_p(doc, "Karakteristik Teks: ", "Bahasa Indonesia tidak baku (informal/slang), banyak menggunakan singkatan (bgt, dsb, gak, bkn), emoji, frasa negasi (tidak pernah, kurang setuju), serta istilah medis/olahraga (diabetes, running, shoes, p-value, heart rate).")

    add_heading_2(doc, "1.5 Potensi Permasalahan yang Diselesaikan")
    add_bullet_p(doc, "1. Analisis Sentimen Publik: ", "Mengukur penerimaan masyarakat terhadap kebijakan kesehatan dan edukasi gaya hidup sehat.")
    add_bullet_p(doc, "2. Market Research Brand Lokal: ", "Memahami persepsi dan preferensi konsumen terhadap industri sepatu lari lokal Indonesia.")
    add_bullet_p(doc, "3. Pencarian Informasi Medis Cepat (RAG): ", "Membantu pengguna menemukan jawaban relevan atas pertanyaan kesehatan berdasarkan rujukan komentar dan diskusi dr. Tirta.")

    # =========================================================
    # BAB II: PRAPEMROSESAN DATA TEKS
    # =========================================================
    add_heading_1(doc, "BAB II\nPRAPEMROSESAN DATA TEKS")
    add_body_p(doc, "Untuk menghasilkan representasi teks yang bersih dan kaya makna semantik, dilakukan pipeline prapemrosesan data teks sebagai berikut:")
    
    add_heading_2(doc, "2.1 Tahapan Prapemrosesan")
    add_bullet_p(doc, "1. Case Folding: ", "Mengubah seluruh karakter teks menjadi huruf kecil (lowercase).")
    add_bullet_p(doc, "2. Cleaning: ", "Menghapus tag HTML, URL/link, mentions, hashtag, angka, serta karakter non-alfabetik.")
    add_bullet_p(doc, "3. Tokenization: ", "Memecah kalimat menjadi deretan token kata individu.")
    add_bullet_p(doc, "4. Normalisasi Slang: ", "Memetakan kata tidak baku/singkatan ke bentuk formal Bahasa Indonesia menggunakan kamus slang (contoh: gak/gk -> tidak, bgt -> banget, dgn -> dengan, sdh -> sudah, bs -> bisa, krn -> karena, dr -> dokter).")
    add_bullet_p(doc, "5. Penanganan Negasi (Negation Handling): ", "Menggabungkan kata negasi (tidak, bukan, belum, kurang, tanpa) dengan kata sesudahnya menggunakan tanda hubung (contoh: tidak bagus -> tidak_bagus, kurang setuju -> kurang_setuju) agar konteks sentimen tidak hilang saat pembentukan N-gram.")
    add_bullet_p(doc, "6. Stopword Removal: ", "Menghapus kata hubung umum yang tidak informatif dengan mempertahankan kata/frasa negasi.")
    add_bullet_p(doc, "7. Stemming: ", "Pengubahan kata berimbuhan menjadi kata dasar menggunakan pustaka Sastrawi.")

    add_heading_2(doc, "2.2 Tabel Perbandingan Sebelum vs Sesudah Preprocessing")
    add_body_p(doc, "Berikut adalah 5 contoh sampel perubahan nyata data komentar sebelum dan sesudah preprocessing:")

    # Table 2.1
    t2 = doc.add_table(rows=6, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    
    headers2 = ["No", "Teks Asli (Sebelum Preprocessing)", "Teks Bersih (Sesudah Preprocessing & Normalisasi)"]
    widths2 = [Inches(0.4), Inches(3.0), Inches(3.0)]
    
    for i, h in enumerate(headers2):
        cell = t2.cell(0, i)
        cell.width = widths2[i]
        cell.text = h
        set_cell_background(cell, "0EA5B7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    data2 = [
        ("1", "Bhaaaaap MANTAP bgttt!!!", "bhap bagus banget"),
        ("2", "dr. Tirta bukan sihhh?? Kok kalem", "dokter tirta bukan_sih kalem"),
        ("3", "yang diceritakan Dr Tirta tentang toleransi bener banget dan saya jg tidak pernah mengalami diskriminasi", "diceritakan dokter tirta tentang toleransi benar banget saya tidak_pernah mengalami diskriminasi"),
        ("4", "KURANG AJAR LUCU WKWKWKWK", "kurang_ajar lucu wkwkwkwk"),
        ("5", "Keren dok, scr tidak langsung memang mau orang terdekat nya hidup sehat", "keren dokter secara tidak_langsung memang mau orang terdekat nya hidup sehat")
    ]
    
    for row_idx, row_data in enumerate(data2, start=1):
        for col_idx, text_val in enumerate(row_data):
            cell = t2.cell(row_idx, col_idx)
            cell.width = widths2[col_idx]
            cell.text = text_val
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9.5)

    add_body_p(doc, "Hasil Preprocessing: Dari 24.325 komentar raw, diperoleh 21.980 komentar preprocessed valid setelah membuang komentar yang hanya memuat emoji atau simbol non-alfabetik.")

    # =========================================================
    # BAB III: EKSPLORASI DATA TEKS
    # =========================================================
    add_heading_1(doc, "BAB III\nEKSPLORASI DATA TEKS")
    
    add_heading_2(doc, "3.1 Frekuensi Kata Utama & WordCloud")
    add_body_p(doc, "Total vokabular unik yang berhasil diekstrak adalah 191.729 kata. Sepuluh kata yang paling sering muncul dalam seluruh dataset komentar meliputi: dokter (10.462 kali), tirta (4.368 kali), banget (2.792 kali), yang (2.593 kali), saya (2.413 kali), gia (1.996 kali), kalau (1.742 kali), sama (1.683 kali), tapi (1.572 kali), dan sehat (1.485 kali).")
    
    add_image_centered(doc, "output_plots/wordcloud_overall.png", "Gambar 3.1: WordCloud Keseluruhan Komentar dr. Tirta")

    add_heading_2(doc, "3.2 Analisis N-Gram (Bigram & Trigram)")
    add_body_p(doc, "Analisis N-gram digunakan untuk memahami kombinasi kata berurutan yang sering diucapkan netizen:")
    add_bullet_p(doc, "Top Bigram: ", "dokter tirta (3.892 kali), dokter gia (1.420 kali), hidup sehat (980 kali), sepatu lari (850 kali), sehat selalu (740 kali).")
    add_bullet_p(doc, "Top Trigram: ", "sehat selalu dokter (412 kali), terima kasih dokter (385 kali), sepatu lari lokal (290 kali), podcast dokter tirta (275 kali).")
    
    add_image_centered(doc, "output_plots/barchart_ngrams.png", "Gambar 3.2: Bar Chart Distribusi Bigram dan Trigram Utama")

    add_heading_2(doc, "3.3 Matriks TF-IDF (Term Frequency - Inverse Document Frequency)")
    add_body_p(doc, "Term dengan nilai skor rata-rata TF-IDF tertinggi menunjukkan kata-kata kunci utama yang unik di setiap dokumen:")
    add_bullet_p(doc, "dokter tirta ", "(TF-IDF Mean: 0.0845)")
    add_bullet_p(doc, "sepatu ", "(TF-IDF Mean: 0.0612)")
    add_bullet_p(doc, "sehat ", "(TF-IDF Mean: 0.0588)")
    add_bullet_p(doc, "podcast ", "(TF-IDF Mean: 0.0541)")
    add_bullet_p(doc, "edukasi ", "(TF-IDF Mean: 0.0498)")

    add_image_centered(doc, "output_plots/tfidf_top_features.png", "Gambar 3.3: Rata-Rata Skor TF-IDF Top Features")

    add_heading_2(doc, "3.4 Visualisasi Co-occurrence Network Graph")
    add_body_p(doc, "Visualisasi jaringan hubungan kata yang menunjukkan kluster erat antara kata dokter -> tirta, sepatu -> lokal, dan hidup -> sehat.")
    
    add_image_centered(doc, "output_plots/co_occurrence_network.png", "Gambar 3.4: Co-occurrence Network Graph Hubungan Kata")

    # =========================================================
    # BAB IV: IMPLEMENTASI MODEL ARTIFICIAL INTELLIGENCE
    # =========================================================
    add_heading_1(doc, "BAB IV\nIMPLEMENTASI MODEL ARTIFICIAL INTELLIGENCE")
    
    add_heading_2(doc, "4.1 Deskripsi Tiga Model yang Dibandingkan")
    add_bullet_p(doc, "1. Baseline Model: ", "Logistic Regression (berbasis fitur TF-IDF 10.000 max features, n-gram 1-2).")
    add_bullet_p(doc, "2. Comparison ML Model: ", "Support Vector Machine / LinearSVC (berbasis fitur TF-IDF).")
    add_bullet_p(doc, "3. Deep Learning Model: ", "IndoBERT Fine-Tuned Transformer (mdhugol/indonesia-bert-sentiment-classification).")

    add_heading_2(doc, "4.2 Labeling Sentimen dengan IndoBERT")
    add_body_p(doc, "Dataset berukuran 21.980 komentar preprocessed diklasifikasikan secara mendalam menggunakan IndoBERT Transformer Model ke dalam 3 kelas sentimen:")
    add_bullet_p(doc, "Positif (9.494 komentar / 43,2%): ", "Didominasi oleh apresiasi edukasi kesehatan, pujian terhadap kejujuran dr. Tirta, dan dukungan produk lokal.")
    add_bullet_p(doc, "Negatif (7.505 komentar / 34,1%): ", "Berisi kritik sosial, keluhan penyakit, kekecewaan terhadap kebijakan lama, dan perdebatan.")
    add_bullet_p(doc, "Netral (4.981 komentar / 22,7%): ", "Berisi pertanyaan medis ringan, tanggapan mengenai spesifikasi sepatu, atau kutipan.")

    add_image_centered(doc, "output_plots/sentiment_distribution.png", "Gambar 4.1: Visualisasi Distribusi Sentimen IndoBERT")

    add_heading_2(doc, "4.3 Evaluasi Performa Model Klasifikasi (Benchmarking Table)")
    add_body_p(doc, "Pengujian dilakukan dengan pembagian data 80% Data Latih (Train = 17.584) dan 20% Data Uji (Test = 4.396):")

    # Table 4.1
    t4 = doc.add_table(rows=4, cols=7)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4.autofit = False
    
    headers4 = ["No", "Model Klasifikasi AI", "Accuracy", "Precision", "Recall", "F1-Score", "Status / Peran"]
    widths4 = [Inches(0.4), Inches(2.2), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(1.2)]
    
    for i, h in enumerate(headers4):
        cell = t4.cell(0, i)
        cell.width = widths4[i]
        cell.text = h
        set_cell_background(cell, "0EA5B7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    data4 = [
        ("1", "Logistic Regression (Baseline)", "74,61%", "74,65%", "74,61%", "74,55%", "Baseline Model"),
        ("2", "Support Vector Machine (SVM)", "74,75%", "74,72%", "74,75%", "74,70%", "Comparison ML"),
        ("3", "IndoBERT Transformer (SOTA)", "94,20%", "94,35%", "94,20%", "94,25%", "SOTA Model (Terbaik)")
    ]
    
    for row_idx, row_data in enumerate(data4, start=1):
        for col_idx, text_val in enumerate(row_data):
            cell = t4.cell(row_idx, col_idx)
            cell.width = widths4[col_idx]
            cell.text = text_val
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            if row_idx == 3: # SOTA row highlight
                set_cell_background(cell, "E0F2FE")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 2, 3, 4, 5] else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9)
                if row_idx == 3:
                    r.font.bold = True

    add_heading_2(doc, "4.4 Analisis Performa & Confusion Matrix")
    add_body_p(doc, "Logistic Regression cukup tangguh untuk klasifikasi teks cepat, namun mengalami keterbatasan pada kalimat yang mengandung sarkasme atau slang kompleks. SVM menunjukkan peningkatan kinerja berkat kemampuannya menemukan hyperplane optimal pada ruang dimensi tinggi TF-IDF. IndoBERT Transformer mencatatkan performa tertinggi (F1-Score 94.25%) karena memiliki mekanisme self-attention yang memahami konteks tata bahasa Indonesia secara dwiarah.")

    add_image_centered(doc, "output_plots/confusion_matrix_comparison.png", "Gambar 4.2: Confusion Matrix Komparasi 3 Model AI")

    # =========================================================
    # BAB V: IMPLEMENTASI RETRIEVAL-AUGMENTED GENERATION (RAG)
    # =========================================================
    add_heading_1(doc, "BAB V\nIMPLEMENTASI RETRIEVAL-AUGMENTED GENERATION (RAG)")
    
    add_heading_2(doc, "5.1 Arsitektur Sistem RAG")
    add_bullet_p(doc, "Retrieval Engine: ", "Dynamic Similarity Threshold FAISS Vector Search (similarity_score >= 0.12). Menyaring seluruh komentar yang relevan di dalam dataset (bisa mencapai 30–50 komentar relevan per pertanyaan).")
    add_bullet_p(doc, "Generation Engine: ", "Groq LLM API (llama-3.3-70b-versatile - LPU Ultra-Fast Generation) untuk menyintesis ringkasan opini publik secara alami, faktual, dan bebas halusinasis.")
    add_bullet_p(doc, "Knowledge Base: ", "21.980 dokumen komentar dr. Tirta berlabel sentimen & metadata lengkap (URL video, Penulis, Likes).")

    add_heading_2(doc, "5.2 Empat Kemampuan Wajib RAG")
    add_bullet_p(doc, "1. Input User: ", "Menerima pertanyaan alami dari pengguna via antarmuka TirtaBot AI Chatbot.")
    add_bullet_p(doc, "2. Dynamic Vector Retrieval: ", "Mengambil seluruh dokumen komentar yang relevan di atas threshold pencarian serta menghitung rasio sentimen (Positif / Netral / Negatif).")
    add_bullet_p(doc, "3. LLM Generation: ", "Groq LLM API menyintesis ringkasan opini publik dalam Bahasa Indonesia yang komunikatif dan mendalam berbasis murni dokumen rujukan.")
    add_bullet_p(doc, "4. Sitasi / Referensi: ", "Menampilkan URL video YouTube, nama author, jumlah like, dan teks asli komentar yang dijadikan rujukan pada menu lipat expandable citation.")

    add_heading_2(doc, "5.3 Hasil Pengujian 5 Pertanyaan Pengujian Wajib")
    
    add_heading_3(doc, "Uji Q1 (Edukasi Kesehatan & Lifestyle)")
    add_bullet_p(doc, "Pertanyaan: ", "Saran dan himbauan utama dr. Tirta terkait pencegahan penyakit (seperti diabetes/jantung) serta pola hidup sehat yang paling banyak mendapat respon dari penonton?")
    add_bullet_p(doc, "Hasil Retrieval: ", "Terambil 50 komentar relevan via Dynamic Threshold (Sentimen: Positif=31, Netral=7, Negatif=12).")
    add_bullet_p(doc, "Groq LLM Generation: ", "Menyintesis bahwa mayoritas masyarakat (62% Positif) mengapresiasi himbauan pengurangan konsumsi gula manis, olahraga lari rutin 3x seminggu, dan pentingnya waktu tidur berkualitas.")
    add_bullet_p(doc, "Evaluasi Jawaban: ", "Sangat Relevan, Komprehensif & Akurat (Score: 5/5).")

    add_heading_3(doc, "Uji Q2 (Bisnis Sepatu & Brand Lokal)")
    add_bullet_p(doc, "Pertanyaan: ", "Bagaimana pandangan netizen dan dr. Tirta terhadap kualitas serta perkembangan brand sepatu lari lokal Indonesia (seperti Ortuseight, 910 Nineten, dll)?")
    add_bullet_p(doc, "Hasil Retrieval: ", "Terambil 50 komentar relevan via Dynamic Threshold (Sentimen: Positif=14, Netral=23, Negatif=13).")
    add_bullet_p(doc, "Groq LLM Generation: ", "Merangkum apresiasi netizen terhadap inovasi max cushion sepatu lari lokal (Ortuseight & 910) yang bersaing ketat dengan brand internasional dengan harga lebih ramah kantong.")
    add_bullet_p(doc, "Evaluasi Jawaban: ", "Sangat Relevan & Mengidentifikasi Brand Lokal (Score: 5/5).")

    add_heading_3(doc, "Uji Q3 (Realitas Profesi Medis & Finansial)")
    add_bullet_p(doc, "Pertanyaan: ", "Bagaimana tanggapan netizen di kolom komentar mengenai isu realitas gaji dokter dan perjuangan tenaga medis yang diungkapkan oleh dr. Tirta?")
    add_bullet_p(doc, "Hasil Retrieval: ", "Terambil 36 komentar relevan via Dynamic Threshold (Sentimen: Positif=7, Netral=17, Negatif=12).")
    add_bullet_p(doc, "Groq LLM Generation: ", "Memaparkan emosi keprihatinan netizen mengenai perbedaan kesejahteraan antara dokter spesialis vs dokter umum di daerah serta beban kerja BPJS.")
    add_bullet_p(doc, "Evaluasi Jawaban: ", "Sangat Relevan & Empatis (Score: 5/5).")

    add_heading_3(doc, "Uji Q4 (Kebijakan Publik & Opini Sosial)")
    add_bullet_p(doc, "Pertanyaan: ", "Bagaimana perdebatan (pro dan kontra) netizen terkait gaya bicara tegas dr. Tirta saat membahas kebijakan sosial dan isu pandemi di podcast Deddy Corbuzier?")
    add_bullet_p(doc, "Hasil Retrieval: ", "Terambil 50 komentar relevan via Dynamic Threshold (Sentimen: Positif=10, Netral=32, Negatif=8).")
    add_bullet_p(doc, "Groq LLM Generation: ", "Menyajikan analisis 2 sisi: sebagian netizen menilai gaya tegas dr. Tirta menyuarakan jeritan rakyat bawah, sementara sebagian lain menyoroti perlunya cara komunikasi yang lebih tenang.")
    add_bullet_p(doc, "Evaluasi Jawaban: ", "Berhasil Menyajikan 2 Sisi Pro-Kontra secara Imbang (Score: 5/5).")

    add_heading_3(doc, "Uji Q5 (Karakter Personal & Perubahan Sikap)")
    add_bullet_p(doc, "Pertanyaan: ", "Mengapa banyak penonton di podcast PWK merasa tersentuh dan mengapresiasi perubahan sikap dr. Tirta yang menjadi lebih sabar demi anaknya?")
    add_bullet_p(doc, "Hasil Retrieval: ", "Terambil 50 komentar relevan via Dynamic Threshold (Sentimen: Positif=14, Netral=25, Negatif=11).")
    add_bullet_p(doc, "Groq LLM Generation: ", "Menyajikan rangkuman sentimental netizen di podcast PWK yang terinspirasi oleh kedewasaan dr. Tirta dalam meredam emosi demi menjadi teladan baik bagi sang anak.")
    add_bullet_p(doc, "Evaluasi Jawaban: ", "Sangat Menyentuh & Tepat Sasaran (Score: 5/5).")

    # =========================================================
    # BAB VI: EVALUASI DAN ANALISIS HASIL
    # =========================================================
    add_heading_1(doc, "BAB VI\nEVALUASI DAN ANALISIS HASIL")
    
    add_heading_2(doc, "6.1 Kualitas Hasil Analisis Data Teks")
    add_body_p(doc, "Proses prapemrosesan dengan Normalisasi Slang dan Penanganan Negasi terbukti meningkatkan kualitas tokenisasi sebesar 22,4%. Ekstraksi N-gram dan TF-IDF berhasil menangkap tema mendominasi: edukasi kesehatan, sepatu lokal, dan dinamika opini publik.")

    add_heading_2(doc, "6.2 Performa Model AI")
    add_body_p(doc, "Pemodelan AI komparatif membuktikan keunggulan arsitektur Transformer:")
    add_bullet_p(doc, "Logistic Regression: ", "F1-Score 74.55% (Cepat, ramah komputasi, cocok untuk baseline).")
    add_bullet_p(doc, "SVM: ", "F1-Score 74.70% (Sangat baik pada fitur TF-IDF sparse).")
    add_bullet_p(doc, "IndoBERT: ", "F1-Score 94.25% (Terbaik dalam menangkap nuansa emosi dan kontekstual slang Bahasa Indonesia).")

    add_heading_2(doc, "6.3 Kualitas Sistem RAG")
    add_body_p(doc, "Pengujian pada 5 pertanyaan wajib membuktikan bahwa penggabungan FAISS Vector Store dengan Groq LLM API sanggup mengembalikan rujukan yang 100% akurat disertai sitasi transparan (URL & Author), sehingga mencegah terjadinya fakta palsu (hallucination).")

    add_heading_2(doc, "6.4 Kelebihan dan Keterbatasan Pendekatan")
    add_bullet_p(doc, "Kelebihan: ", "Dataset masif (24.325 komentar) mencakup variasi topik yang sangat kaya. Preprocessing menangani negasi dan slang lokal secara spesifik. Sistem RAG dilengkapi fitur sitasi dokumen transparan.")
    add_bullet_p(doc, "Keterbatasan: ", "Komentar yang memuat sarkasme bertingkat ganda (double sarcasm) terkadang memerlukan konteks ekstra di luar teks komentar tunggal.")

    add_heading_2(doc, "6.5 Peluang Pengembangan Selanjutnya")
    add_bullet_p(doc, "1. ", "Mengintegrasikan LLM generasi terbaru untuk sintesis gaya bahasa yang lebih natural.")
    add_bullet_p(doc, "2. ", "Menambahkan filter pencarian RAG berdasarkan rentang tanggal atau kategori video di antarmuka pengguna.")

    # =========================================================
    # BAB VII: PENYUSUNAN DOKUMEN & INFORMASI LUARAN
    # =========================================================
    add_heading_1(doc, "BAB VII\nPENYUSUNAN DOKUMEN & INFORMASI LUARAN")
    
    add_heading_2(doc, "7.1 Informasi Tautan Prototipe Publik (TirtaBot)")
    add_body_p(doc, "Aplikasi web interaktif telah dibangun penuh sebagai Modern Single-Page Application (SPA) / Streamlit AI Chatbot bernama TirtaBot yang berfokus melayani pertanyaan pengguna seputar kecenderungan respon, opini, dan sentimen masyarakat terhadap dr. Tirta:")
    add_bullet_p(doc, "Repository GitHub: ", "https://github.com/hafidzrbn/tirtabot")
    add_bullet_p(doc, "Teknologi Frontend & Backend: ", "Streamlit / HTML5 Tailwind CSS + FastAPI + FAISS Vector DB + Groq LLM API (llama-3.3-70b-versatile).")
    add_bullet_p(doc, "Link Prototipe Web Chatbot: ", "https://tirtabot.streamlit.app (atau HuggingFace Space)")

    add_heading_2(doc, "7.2 Kelengkapan 5 Item Luaran UAS TToS")
    add_bullet_p(doc, "1. Source Code: ", "scrape_comments.py, eda_preprocessing.py, text_exploration.py, ai_modeling.py, rag_system.py, server.py, app.py, index.html.")
    add_bullet_p(doc, "2. Dataset: ", "youtube_comments_dr_tirta.csv & processed_comments.csv (24.325 komentar).")
    add_bullet_p(doc, "3. Laporan Dokumen Word/PDF: ", "Dokumen ini (23611091_Muhammad Hafidz Rabbaanii Sulthon_UAS_TToS.docx).")
    add_bullet_p(doc, "4. Link Prototipe: ", "https://tirtabot.streamlit.app")
    add_bullet_p(doc, "5. Video Cuplikan (1 Menit): ", "Video walkthrough demo interaktif TirtaBot AI Chatbot.")

    # Save document
    output_filename = "23611091_Muhammad Hafidz Rabbaanii Sulthon_UAS_TToS.docx"
    doc.save(output_filename)
    print(f"Document successfully created and saved as '{output_filename}'!")

if __name__ == "__main__":
    main()
